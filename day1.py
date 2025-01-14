from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 1. Add a paragraph
paragraph = doc.add_paragraph("Here is a sample paragraph. ")

# 2. Add a run for bold text
bold_run = paragraph.add_run("This text is bold. ")
bold_run.bold = True

# 3. Add a run for italic text
italic_run = paragraph.add_run("This text is italic. ")
italic_run.italic = True

# 4. Customize style (font size, typeface)
font_run = paragraph.add_run("Customized text! ")
font_run.font.name = "Arial"
font_run.font.size = Pt(14)


#doc.save("Level2_Formatting.docx")


#more pratices:

doc2 = Document()

para1 = doc2.add_paragraph()
run1 = para1.add_run("This is the first paragraph, and it is left-aligned. ")
run1.bold = True
run1.add_text("This part is bold. ")
run1.italic = True
run1.add_text("This part is italic. ")
run1.underline = True
run1.add_text("This part is underlined. ")
para1.alignment = WD_ALIGN_PARAGRAPH.LEFT

para2 = doc2.add_paragraph()
run2 = para2.add_run("This is the second paragraph, and it is center-aligned. ")
run2.bold = True
run2.add_text("Bold.")
run2.italic = True
run2.add_text("Italic. ")
run2.underline = True
run2.add_text("Underlined. ")
para2.alignment = WD_ALIGN_PARAGRAPH.CENTER

para3 = doc2.add_paragraph()
run3 = para3.add_run("This is the third paragraph, and it is right-alignment. ")
run3.add_text("This sentence will have a custom font and size.")
run3.font.name = "Arial"
run3.font.size = Pt(14)
para3.alignment = WD_ALIGN_PARAGRAPH.RIGHT


doc2.save("Practice1.docx")