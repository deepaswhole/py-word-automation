from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

#doc = Document()

#table = doc.add_table(rows=2, cols=3)
#table.style = 'Table Grid'

#hdr_cells = table.rows[0].cells
#hdr_cells[0].text = "Library"
#hdr_cells[1].text = "Purpose"
#hdr_cells[2].text = "Version"

#row_cells = table.rows[1].cells
#row_cells[0].text = "python-docx"
#row_cells[1].text = "Word Automation"
#row_cells[2].text = "0.9.11"

#doc.save("Level3_Tables.docx")

doc = Document()


table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

hdr_row = table.rows[0]
hdr_row.cells[0].text = "Python Library Information"
hdr_row.cells[0].merge(hdr_row.cells[1]).merge(hdr_row.cells[2])

hdr_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

table.cell(1,0).text = "Library Name"
table.cell(1,1).text = "Typical Use-Case"
table.cell(1,2).text = "Personal Rating"

libraries = [
    ["Pandas", "Data manipulation and analysis", "9/10"],
    ["numpy", "Scientific computing and array", "8.5/10"],
    ["matplotlib", "Data visualization", "8/10"],
    ["scikit-learn", "Machine learning", "9/10"],
    ["python-docx", "Word document automation", "8/10"]
]

for i, (name, use_case, rating) in enumerate(libraries, start=2):
    table.cell(i, 0).text = name
    table.cell(i, 1).text = use_case
    table.cell(i, 2).text = rating


doc.save('level3 practice with cells merge.docx')


